from __future__ import annotations

from datetime import date
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
BASE = ROOT / "manuscript" / "submission_package" / "targets" / "journal_of_biomedical_informatics"
TITLE = (
    "A Reproducible Uncertainty-Aware Multi-Omics Network Pipeline for Ovarian Cancer: "
    "Stable Cross-Layer Hubs, External Immune-Context Validation, and Direct CAR-Product Benchmarking"
)
SHORT_TITLE = "Uncertainty-Aware Ovarian Multi-Omics Informatics Pipeline"


def refs() -> list[str]:
    return [
        "The Cancer Genome Atlas Research Network. Integrated genomic analyses of ovarian carcinoma. Nature. 2011;474(7353):609-615. doi:10.1038/nature10166.",
        "National Cancer Institute. Genomic Data Commons (GDC) Data Portal [Internet]. Available from: https://portal.gdc.cancer.gov/.",
        "National Cancer Institute. GDC API Users Guide: Downloading Files [Internet]. Available from: https://docs.gdc.cancer.gov/API/Users_Guide/Downloading_Files/.",
        "Colaprico A, Silva TC, Olsen C, Garofano L, Cava C, Garolini D, et al. TCGAbiolinks: an R/Bioconductor package for integrative analysis of TCGA data. Nucleic Acids Res. 2016;44(8):e71. doi:10.1093/nar/gkv1507.",
        "Argelaguet R, Velten B, Arnol D, Dietrich S, Zenz T, Marioni JC, et al. Multi-Omics Factor Analysis-a framework for unsupervised integration of multi-omics data sets. Mol Syst Biol. 2018;14(6):e8124. doi:10.15252/msb.20178124.",
        "Argelaguet R, Arnol D, Bredikhin D, Deloro Y, Velten B, Marioni JC, et al. MOFA+: a statistical framework for comprehensive integration of multi-modal single-cell data. Genome Biol. 2020;21(1):111. doi:10.1186/s13059-020-02015-1.",
        "Rohart F, Gautier B, Singh A, Le Cao KA. mixOmics: an R package for omics feature selection and multiple data integration. PLoS Comput Biol. 2017;13(11):e1005752. doi:10.1371/journal.pcbi.1005752.",
        "Singh A, Shannon CP, Gautier B, Rohart F, Vacher M, Tebbutt SJ, et al. DIABLO: an integrative approach for identifying key molecular drivers from multi-omics assays. Bioinformatics. 2019;35(17):3055-3062. doi:10.1093/bioinformatics/bty1054.",
        "Cerami E, Gao J, Dogrusoz U, Gross BE, Sumer SO, Aksoy BA, et al. The cBio cancer genomics portal: an open platform for exploring multidimensional cancer genomics data. Cancer Discov. 2012;2(5):401-404. doi:10.1158/2159-8290.CD-12-0095.",
        "Gao J, Aksoy BA, Dogrusoz U, Dresdner G, Gross B, Sumer SO, et al. Integrative analysis of complex cancer genomics and clinical profiles using the cBioPortal. Sci Signal. 2013;6(269):pl1. doi:10.1126/scisignal.2004088.",
        "National Cancer Institute. Proteomic Data Commons (PDC) [Internet]. Available from: https://pdc.cancer.gov/.",
        "Love MI, Huber W, Anders S. Moderated estimation of fold change and dispersion for RNA-seq data with DESeq2. Genome Biol. 2014;15(12):550. doi:10.1186/s13059-014-0550-8.",
        "Hagberg AA, Schult DA, Swart PJ. Exploring network structure, dynamics, and function using NetworkX. In: Proceedings of the 7th Python in Science Conference; 2008. p. 11-15.",
    ]


def set_font(doc: Document, name: str = "Times New Roman", size: int = 12) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)


def heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc: Document, df: pd.DataFrame, title: str, max_rows: int | None = None) -> None:
    doc.add_paragraph(title).runs[0].bold = True
    dfx = df.copy()
    if max_rows is not None:
        dfx = dfx.head(max_rows)
    table = doc.add_table(rows=1, cols=len(dfx.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, c in enumerate(dfx.columns):
        hdr[i].text = str(c)
    for _, row in dfx.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(dfx.columns):
            cells[i].text = str(row[c])
    doc.add_paragraph("")


def add_figure(doc: Document, img_path: Path, caption: str, width: float = 6.1) -> None:
    if not img_path.exists():
        doc.add_paragraph(f"[Missing figure: {img_path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Inches(width))
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def save(doc: Document, path: Path) -> Path:
    try:
        doc.save(path)
        return path
    except PermissionError:
        alt = path.with_name(path.stem + "_updated.docx")
        doc.save(alt)
        return alt


def make_graphical_abstract(out_png: Path, n_core: int, n_prot: int, top_hubs: str) -> None:
    plt.figure(figsize=(14, 8), dpi=220)
    ax = plt.gca()
    ax.axis("off")

    def box(x, y, w, h, txt, fc):
        rect = plt.Rectangle((x, y), w, h, fc=fc, ec="#1f2937", lw=1.4)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, txt, ha="center", va="center", fontsize=10)

    box(0.03, 0.64, 0.24, 0.24, "TCGA-OV matched cohort\nRNA + CNA + methylation + mutation\nOptional protein", "#E9F2FF")
    box(0.37, 0.64, 0.24, 0.24, "MOFA-like + DIABLO-like integration\nnetwork centrality\nperturbation + sensitivity", "#EEF9EF")
    box(0.71, 0.64, 0.24, 0.24, "External support\novarian TIL validation\nCAR-product raw-read benchmark", "#FFF4E5")
    box(
        0.03,
        0.16,
        0.92,
        0.30,
        f"Core matched cohort n={n_core} | Protein-matched n={n_prot}\nStable hubs: {top_hubs}\nMain contribution: reproducible mechanism-oriented evidence with external validation assets",
        "#FFFFFF",
    )
    ax.annotate("", xy=(0.35, 0.76), xytext=(0.27, 0.76), arrowprops=dict(arrowstyle="->", lw=2))
    ax.annotate("", xy=(0.69, 0.76), xytext=(0.61, 0.76), arrowprops=dict(arrowstyle="->", lw=2))
    ax.text(0.5, 0.98, "Graphical Abstract: Reproducible Ovarian Multi-Omics Informatics Pipeline", ha="center", va="top", fontsize=14, fontweight="bold")
    plt.tight_layout()
    out_png.parent.mkdir(parents=True, exist_ok=True)
    plt.savefig(out_png, bbox_inches="tight")
    plt.close()


def build_main() -> Path:
    tables = ROOT / "results" / "tables"
    figs = ROOT / "results" / "figures"
    nets = ROOT / "results" / "networks"
    sample = pd.read_csv(tables / "sample_matching_summary.csv")
    bench = pd.read_csv(tables / "model_benchmark.csv")
    advanced_ml = pd.read_csv(tables / "advanced_ml_benchmark.csv")
    perm = pd.read_csv(tables / "permutation_test_auc.csv")
    motif = pd.read_csv(tables / "cart_motif_benchmark.csv") if (tables / "cart_motif_benchmark.csv").exists() else pd.DataFrame()
    readiness = pd.read_csv(tables / "cart_reference_alignment_readiness.csv") if (tables / "cart_reference_alignment_readiness.csv").exists() else pd.DataFrame()
    hub_stab = pd.read_csv(nets / "network_centrality_stability.csv")
    ext_immune = pd.read_csv(tables / "external_ovarian_immune_summary.csv")
    cart_qc = pd.read_csv(tables / "cart_direct_benchmark_qc.csv")

    n_core = int(sample.loc[sample["metric"] == "n_patients_intersection_all_main_layers", "value"].iloc[0])
    n_prot = int(sample.loc[sample["metric"] == "n_patients_protein", "value"].iloc[0])
    top_auc = bench.sort_values("auc", ascending=False).iloc[0]
    top_cox = bench.sort_values("cox_c_index", ascending=False).iloc[0]
    top_ml = advanced_ml.sort_values("auc", ascending=False).iloc[0]
    perm_p = float(perm["p_value_right_tail"].iloc[0]) if not perm.empty else float("nan")
    top_hubs = ", ".join(hub_stab.head(4)["node"].tolist())

    ga_path = BASE / "JBI_Graphical_Abstract.png"
    make_graphical_abstract(ga_path, n_core, n_prot, top_hubs)

    doc = Document()
    set_font(doc)

    heading(doc, TITLE, 1)
    p = doc.add_paragraph("Dr Siddalingaiah H S, Professor, Community Medicine, Shridevi Institute of Medical Sciences and Research Hospital, Tumkur, India. ")
    p.add_run("Email: hssling@yahoo.com | Phone: 8941087719 | ORCID: 0000-0002-4771-8285")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading(doc, "Abstract", 2)
    abstract = (
        f"Objective: To build a reproducible multi-omics informatics pipeline that distinguishes robust mechanistic signals from unstable predictive claims in TCGA ovarian cancer. "
        f"Materials and methods: Public TCGA-OV mutation, copy-number alteration, methylation, RNA, clinical, and optional protein data were harmonized at patient level and analyzed using latent integration, supervised multiblock modeling, multilayer network centrality, perturbation, sensitivity analysis, ablation, and advanced machine-learning calibration. External support was added through an ovarian tumor-infiltrating CD8 T-cell RNA-seq dataset and a direct CAR-product raw-read benchmark dataset. "
        f"Results: The core matched cohort included {n_core} patients; the processed protein-matched subset included {n_prot}. RNA modules achieved the highest all-available AUC ({top_auc['auc']:.3f}, 95% CI {top_auc['auc_ci_low']:.3f}-{top_auc['auc_ci_high']:.3f}), whereas CNA achieved the highest Cox C-index ({top_cox['cox_c_index']:.3f}, 95% CI {top_cox['cox_c_index_ci_low']:.3f}-{top_cox['cox_c_index_ci_high']:.3f}). Stable hubs were {top_hubs}. External ovarian validation showed higher exhaustion-regulatory scores in CD137-positive and CD137-negative PD-1high CD39+ T-cell states than in PD-1+ CD39- controls. Direct CAR benchmarking successfully ingested paired FASTQ files from a public CD22 CAR-T product run and generated sequence-QC outputs. "
        "Discussion and conclusion: The pipeline’s strongest contribution is a reproducible evidence structure linking uncertainty-aware network biology, external ovarian immune-context support, and a direct raw-read CAR benchmark without overstating construct-recovery claims."
    )
    doc.add_paragraph(abstract)
    doc.add_paragraph("Keywords: ovarian cancer; multi-omics; network biology; translational bioinformatics; reproducible research; CAR-T benchmarking")

    heading(doc, "1. Introduction", 2)
    doc.add_paragraph(
        "A central problem in translational bioinformatics is how to distinguish computationally interesting patterns from findings that remain stable when examined through complementary analytical views. This is especially relevant in multi-omics cancer studies, where feature dimensionality is high, matched patient cohorts are modest, and apparent gains from integration can be driven by unstable estimates rather than reproducible biology [1-11]."
    )
    doc.add_paragraph(
        "We therefore reframed the ovarian cancer question as an informatics problem: can a public-data pipeline recover cross-layer signals that remain coherent across latent integration, supervised discrimination, network centrality, perturbation analysis, external immune-context comparison, and raw-read benchmark extensions? The aim was not merely to optimize a classifier but to produce a reusable and auditable evidence structure."
    )
    doc.add_paragraph(
        "This distinction matters because public cancer resources are increasingly used to support translational claims, yet the analytical narrative often collapses several different forms of evidence into a single headline metric. A modest AUC, a stable network hub, an interpretable pathway module, and an externally coherent immune-state signal do not carry the same meaning. A stronger computational manuscript should therefore separate them, quantify them, and state explicitly which claims are supported by which data products."
    )
    doc.add_paragraph(
        "Ovarian cancer is a suitable test case for this strategy. TCGA-OV offers broad multi-layer public coverage, but the effective matched cohort shrinks as data types are intersected, especially when protein is included. That makes it a realistic setting in which uncertainty-aware reporting is not optional. If a method remains useful under these conditions, it is more likely to be informative in other public multi-omics cohorts as well."
    )
    doc.add_paragraph(
        "The present study was therefore designed around traceability as much as around biological discovery. Every major result is linked to an inspectable intermediate artifact: matched-sample counts, per-layer features, benchmark intervals, network stability summaries, perturbation deltas, external validation tables, or sequence-file QC products. This reporting structure serves both oncology readers interested in ovarian biology and biomedical informatics readers interested in how evidence travels from raw public files to manuscript-ready claims."
    )
    doc.add_paragraph(
        "A further motivation was the recurring mismatch between what public datasets can answer and what readers hope they can answer. TCGA-OV can support integrated systems-level analyses of tumor biology, but it does not directly resolve engineered therapeutic sequence architecture. External immune datasets can strengthen a biological direction, but they do not automatically validate a risk classifier. By building these distinctions directly into the pipeline, the study models a more careful standard for computational oncology reporting."
    )
    doc.add_paragraph(
        "The resulting manuscript is positioned as a pipeline-and-evidence paper rather than a pure performance paper. Its scientific value depends on whether the workflow can recover coherent ovarian biology while also making the limits of public inference explicit. That combination is more informative than a superficially stronger but less defensible predictive narrative."
    )

    heading(doc, "2. Materials and methods", 2)
    doc.add_paragraph(
        "TCGA-OV data were retrieved through GDC-oriented workflows and harmonized at patient level with strict ID matching [2-4]. The main analysis required matched mutation, CNA, methylation, and RNA data; protein was evaluated in a fairness-restricted subset. Latent cross-omics structure was modeled using MOFA-like factors [5,6], supervised multiblock structure using a DIABLO-like framework [7,8], and downstream relations through a multilayer NetworkX-based graph with perturbation and sensitivity analysis [13]."
    )
    doc.add_paragraph(
        "To strengthen external support, we added two extensions. First, an ovarian tumor-infiltrating CD8 T-cell RNA-seq dataset (GSE160705/SRX9423612) was used for immune-context comparison using gene-set-derived proxy scores. Second, a public CD22 CAR-T product RNA-seq run (SRR31001810) was downloaded from ENA/SRA and passed through a lightweight raw-read ingestion and QC benchmark. This second extension was designed to validate the external raw-read branch and dataset suitability, not to claim CAR construct recovery without a validated reference panel."
    )
    doc.add_paragraph(
        "Preprocessing was kept strict and fully script-driven. Raw GDC downloads remained immutable, feature construction was derived from processed layer-specific matrices, and patient matching summaries were written at every major stage. RNA and protein were represented both as normalized high-dimensional matrices and lower-dimensional module summaries; CNA, methylation, and mutation were also summarized into burden-style features to support block-wise comparisons and input-output ablation experiments."
    )
    doc.add_paragraph(
        "The network layer linked RNA modules, latent factors, protein modules, and phenotype-relevant nodes. Hub ranking combined centrality metrics and bootstrap resampling to reduce sensitivity to any single graph realization. Perturbation analysis then measured how strongly the global PageRank distribution changed when high-priority hubs were dampened. Sensitivity curves across a perturbation-fraction grid were used to identify whether hub influence was monotonic and graded rather than threshold-dependent."
    )
    doc.add_paragraph(
        "The external ovarian validation branch was analyzed at expression level using the normalized GEO workbook and group labels recovered from the GEO series metadata. The direct CAR benchmark branch was intentionally lightweight: paired FASTQ files were downloaded, checksummed, and summarized by read-length and nucleotide composition from sampled reads. This provided a reproducible benchmark for public raw-read handling while avoiding unsupported claims of CAR construct recovery."
    )
    doc.add_paragraph(
        "All downstream outputs, including inventories, checksums, figures, and manuscript-ready tables, were written into the project workspace so that both scientific interpretation and computational provenance could be audited together."
    )
    doc.add_paragraph(
        "Model benchmarking used repeated resampling to report interval estimates rather than isolated point estimates wherever possible. The same principle guided hub prioritization: nodes were ranked not only by centrality magnitude but also by bootstrap stability and top-k retention frequency. For the machine-learning branch, the emphasis was on calibrated interpretation of limited predictive signal rather than forced performance inflation."
    )
    doc.add_paragraph(
        "The project also included explicit scope controls. Synthetic data fallbacks were removed from QC and parsing stages so that all reported outputs derive from true public files. The CAR-related branch was restricted to what the downloaded data can justify directly: file-level benchmarking, sampled-read QC, heuristic motif screening, and readiness planning for future alignment against a validated external reference panel."
    )

    heading(doc, "3. Results", 2)
    doc.add_paragraph(
        f"The core matched TCGA-OV cohort included {n_core} patients and the processed protein-matched subset included {n_prot}. RNA modules produced the highest all-available AUC ({top_auc['auc']:.3f}, 95% CI {top_auc['auc_ci_low']:.3f}-{top_auc['auc_ci_high']:.3f}), whereas CNA produced the highest Cox C-index ({top_cox['cox_c_index']:.3f}, 95% CI {top_cox['cox_c_index_ci_low']:.3f}-{top_cox['cox_c_index_ci_high']:.3f}). Stable hub ranking consistently prioritized {top_hubs}."
    )
    doc.add_paragraph(
        "These benchmark patterns reinforced an important methodological point: the layer with the strongest direct discriminative signal was not necessarily the layer with the strongest survival-ordering information, and neither of those was identical to the layer structure that dominated the integrated network. In other words, classification, ranking, and mechanism-oriented organization reflected different aspects of the same matched cohort."
    )
    doc.add_paragraph(
        "The latent hubs were especially informative because they persisted across several analytical stress tests. They ranked highly in bootstrap centrality summaries, drove the largest perturbation responses, and occupied the most central positions in the directed layer graph. This persistence argues against interpreting them as artifacts of one fitting procedure and instead supports their use as integrated regulatory states for downstream follow-up."
    )
    doc.add_paragraph(
        "External ovarian immune-context validation reinforced the biological plausibility of the project’s immune-state axis. In GSE160705, exhaustion-regulatory scores were materially higher in CD137-positive and CD137-negative PD-1high CD39+ CD8 T-cell groups than in PD-1+ CD39- control cells. This does not validate the TCGA-OV risk model directly, but it supports the relevance of the immune-state direction captured by the RNA-derived proxy framework."
    )
    doc.add_paragraph(
        "The same external dataset also showed coherent differences for T-cell core, cytolytic, and interferon/antigen-presentation score families across the three annotated groups. Because this dataset was collected in a different biological context from the main TCGA cohort, it is best interpreted as directional support for the immune-state framework rather than as a direct replication of the outcome model. That distinction is analytically important and is preserved explicitly in the pipeline outputs."
    )
    doc.add_paragraph(
        "The direct CAR benchmark confirmed the practicality of the external raw-read branch. The CD22 CAR-T product run SRR31001810 was downloaded successfully as paired FASTQ files and yielded stable lightweight QC summaries with mean sampled read lengths of approximately 75.5 bases and GC fractions of 0.48 to 0.49. These results establish a real public raw-read benchmark for future alignment-based screening once a validated reference panel is available."
    )
    doc.add_paragraph(
        "This branch improves the project in a different way from the external ovarian validation branch. It does not add ovarian biology, but it proves that the workflow can retrieve, inventory, checksum, and summarize true CAR-product sequence data in a reproducible manner. That is a necessary intermediate step before any alignment-based CAR screening claim can be made responsibly."
    )
    doc.add_paragraph(
        "The conservative motif benchmark further clarified this boundary. Across sampled reads from both mates, the heuristic screen did not identify the selected backbone-like motifs at detectable frequency. This negative result should not be overinterpreted: short-read orientation, transcript design, sequencing depth distribution, and construct-specific architecture all limit motif-based inference. Its practical value is to show why a validated alignment panel is required before any sequence-level conclusion is reported."
    )
    doc.add_paragraph(
        f"Advanced machine-learning analyses did not materially overturn the uncertainty-aware picture. The best flexible model achieved AUC {top_ml['auc']:.3f}, while permutation testing remained non-significant at p={perm_p:.4f}. In this dataset, the most defensible outputs are therefore stable hubs, pathway-level directions, and uncertainty-aware comparative benchmarks rather than aggressive claims of high clinical prediction accuracy."
    )
    doc.add_paragraph(
        "A dedicated alignment-readiness audit translated the CAR requirement into operational terms. The workspace now includes a reference-panel staging area, readiness checks for the expected FASTA and metadata files, and a shell-command template for future alignment. At the time of manuscript assembly, the benchmark branch remained appropriately blocked at the construct-inference stage because no validated public reference panel had been installed."
    )
    doc.add_paragraph(
        "Input-output ablation analyses also showed that integrated performance gains were selective rather than universal. Some outputs benefited most from RNA module summaries, whereas others retained stronger ordering information from CNA-derived features. This heterogeneity supports the study's underlying design choice: a multi-layer ovarian model should preserve layer-specific interpretability instead of collapsing all views into a single opaque score."
    )

    add_table(
        doc,
        bench[["model", "auc", "auc_ci_low", "auc_ci_high", "cox_c_index", "cox_c_index_ci_low", "cox_c_index_ci_high"]],
        "Table 1. Core TCGA-OV benchmark summary.",
    )
    add_table(
        doc,
        hub_stab[["node", "rank_score_mean", "rank_score_ci_low", "rank_score_ci_high", "topk_freq"]].head(6),
        "Table 2. Stable cross-layer hub summary.",
    )
    add_table(
        doc,
        ext_immune,
        "Table 3. External ovarian immune-context validation summary.",
    )
    add_table(
        doc,
        cart_qc,
        "Table 4. Direct CAR-product raw-read benchmark QC summary.",
    )
    if not motif.empty:
        add_table(
            doc,
            motif,
            "Table 5. Heuristic CAR motif benchmark summary.",
        )
    if not readiness.empty:
        add_table(
            doc,
            readiness,
            "Table 6. CAR reference-panel alignment readiness audit.",
        )

    add_figure(doc, figs / "multilayer_network_graph.png", "Figure 1. Integrated multilayer network centered on stable latent hubs.")
    add_figure(doc, figs / "perturbation_bootstrap_ci.png", "Figure 2. Perturbation effect sizes with bootstrap confidence intervals.")
    add_figure(doc, figs / "external_ovarian_immune_scores.png", "Figure 3. External ovarian immune-state score distributions across CD137-defined cell groups.")
    add_figure(doc, figs / "external_ovarian_immune_heatmap.png", "Figure 4. Correlation structure of external ovarian immune-state scores.")
    add_figure(doc, ga_path, "Figure 5. Graphical abstract.")

    heading(doc, "4. Discussion", 2)
    doc.add_paragraph(
        "From an informatics perspective, the value of this work lies in how different evidence layers are separated and audited. TCGA-OV supports strong mechanism-oriented network inference, but only modest direct prediction. External ovarian RNA-seq provides biologically relevant immune-context support, whereas direct CAR-product sequencing provides a realistic benchmark for the raw-read branch. Treating those two extensions as equivalent would be methodologically weak; separating them makes the workflow more credible and reusable."
    )
    doc.add_paragraph(
        "The resulting package is stronger than a single-cohort classifier paper. It combines strict matching, uncertainty calibration, hub stability, perturbation analysis, external immune-context comparison, sequence-file inventory, direct FASTQ benchmark acquisition, and manuscript-ready reproducibility assets. In practical terms, it provides both a prioritized biological shortlist and a tested informatics framework for future extension."
    )
    doc.add_paragraph(
        "The study also illustrates a broader principle for translational bioinformatics. Public datasets often tempt analysts to move directly from model performance to biological narrative. In this project, a better sequence was possible: first quantify the stability of network structure, then assess whether the same biological direction remains visible in an external ovarian immune dataset, and only then extend the workflow toward a true raw-read CAR benchmark. Each step answers a different methodological question and reduces the risk of conflating scope."
    )
    doc.add_paragraph(
        "This is particularly relevant for engineered-construct questions. Ovarian tumor datasets do not become CAR-sequence resources simply because they include T-cell populations. Conversely, CAR-product raw-read datasets do not become ovarian validation resources simply because they are immunologically relevant. By partitioning these purposes, the present workflow produces a cleaner and more defensible evidence chain."
    )
    doc.add_paragraph(
        "The network-centric results are also scientifically useful on their own terms. The stable prominence of LF8, LF5, LF6, and LF7 suggests that the most reproducible cross-layer structure in this cohort is concentrated in latent regulatory states rather than in isolated single genes. Because the same hubs remain visible under perturbation and resampling, they form a stronger shortlist for future mechanistic follow-up than any single thresholded association."
    )
    doc.add_paragraph(
        "Equally important, the study makes explicit what it does not show. The available evidence does not justify claims of a ready-to-deploy ovarian clinical decision model, nor does it justify recovery of therapeutic CAR design from the current benchmark branch. These negatives are not weaknesses in the reporting. They are examples of disciplined scope control, which is itself a contribution in a field where public-data analyses can otherwise drift into exaggerated translational framing."
    )
    doc.add_paragraph(
        "The manuscript is therefore intended for a readership that values methodological clarity over inflated certainty. For biomedical informatics journals, that framing is appropriate because reproducibility, auditable provenance, and honest uncertainty are central scientific outcomes rather than mere reporting accessories. The current workflow demonstrates how a public cancer study can be built to preserve those properties from download through manuscript assembly."
    )
    doc.add_paragraph(
        "A second future direction is comparative deployment of the workflow across other TCGA cohorts with similar matched-layer shrinkage. That would allow investigators to compare when integrated network evidence remains stable and when predictive performance collapses under honest validation. Such a comparative informatics resource would likely be more valuable than another isolated case study because it would turn one ovarian analysis into a general benchmark design pattern."
    )
    doc.add_paragraph(
        "For future work, the most natural next step is a validated, alignment-based CAR reference panel evaluated on the downloaded CAR-product dataset. Once that stage is implemented, the existing file inventory, checksums, and QC outputs can act as the audited starting point. In parallel, the external ovarian immune branch can be expanded to additional GEO or SRA cohorts to determine whether the immune-state direction remains stable across independent ovarian microenvironments."
    )

    heading(doc, "5. Limitations", 2)
    doc.add_paragraph(
        "This remains a public-data computational study with no new wet-lab validation. The external ovarian validation dataset is small, and the direct CAR benchmark currently stops at ingestion and QC rather than full alignment-based construct screening. Those constraints are explicit in the pipeline and are preferable to overstated claims."
    )
    doc.add_paragraph(
        "The latent factors are model-derived constructs and should be treated as hypothesis-generating states rather than experimentally validated pathways. The external ovarian validation branch uses processed public counts rather than a unified re-alignment with the TCGA branch. The direct CAR benchmark is intentionally limited by the absence of a validated public reference panel suitable for responsible alignment-based interpretation. None of these limitations invalidate the pipeline, but they define the boundary of the current claim set."
    )

    heading(doc, "6. Conclusion", 2)
    doc.add_paragraph(
        "A reproducible ovarian cancer multi-omics informatics pipeline identified stable cross-layer hubs, confirmed the importance of uncertainty-aware interpretation, and added two complementary external support layers: ovarian immune-context validation and direct CAR-product raw-read benchmarking. This integrated evidence structure is suitable for journal submission as a translational bioinformatics contribution."
    )

    heading(doc, "Declarations", 2)
    doc.add_paragraph("Funding: This research did not receive any specific grant from funding agencies in the public, commercial, or not-for-profit sectors.")
    doc.add_paragraph("Competing interests: The author declares no competing interests.")
    doc.add_paragraph("Ethics: This study used public de-identified datasets and did not require new institutional ethics approval or participant consent.")
    doc.add_paragraph("Data and code availability: Source data are available from GDC, GEO, ENA/SRA, cBioPortal, and PDC under their respective terms. Derived outputs and reproducibility assets are available in this project workspace.")
    doc.add_paragraph(
        "Declaration of generative AI and AI-assisted technologies in the manuscript preparation process: During the preparation of this work, AI-assisted tools were used to support drafting, organization, and language refinement. All scientific content, factual verification, interpretation, and final editing were reviewed by the author, who takes full responsibility for the manuscript."
    )

    heading(doc, "References", 2)
    for i, r in enumerate(refs(), start=1):
        doc.add_paragraph(f"{i}. {r}")

    return save(doc, BASE / "JBI_Main_Manuscript.docx")


def build_supplement() -> Path:
    tables = ROOT / "results" / "tables"
    figs = ROOT / "results" / "figures"
    nets = ROOT / "results" / "networks"

    doc = Document()
    set_font(doc)
    heading(doc, "Journal of Biomedical Informatics Supplementary Appendix", 1)
    doc.add_paragraph(f"Generated: {date.today().isoformat()}")
    doc.add_paragraph("Author: Dr Siddalingaiah H S")

    for path, title, max_rows in [
        (tables / "model_benchmark.csv", "Supplementary Table S1. All-sample benchmark.", None),
        (tables / "model_benchmark_protein_matched.csv", "Supplementary Table S2. Protein-matched benchmark.", None),
        (nets / "network_centrality_stability.csv", "Supplementary Table S3. Full hub stability table.", 20),
        (tables / "perturbation_delta.csv", "Supplementary Table S4. Perturbation delta table.", 20),
        (tables / "advanced_ml_benchmark.csv", "Supplementary Table S5. Advanced ML benchmark.", None),
        (tables / "input_output_ablation_auc.csv", "Supplementary Table S6. Input-output ablation summary.", None),
        (tables / "immune_receptor_proxy_summary.csv", "Supplementary Table S7. TCGA-OV immune proxy summary.", None),
        (tables / "external_ovarian_immune_summary.csv", "Supplementary Table S8. External ovarian immune validation summary.", None),
        (tables / "cart_direct_benchmark_qc.csv", "Supplementary Table S9. Direct CAR-product FASTQ benchmark QC.", None),
        (tables / "car_t_architecture_metadata.csv", "Supplementary Table S10. Public CAR architecture metadata.", 20),
        (tables / "cart_motif_benchmark.csv", "Supplementary Table S11. Heuristic CAR motif benchmark.", None),
        (tables / "cart_reference_alignment_readiness.csv", "Supplementary Table S12. CAR reference alignment readiness audit.", None),
    ]:
        if path.exists():
            add_table(doc, pd.read_csv(path), title, max_rows=max_rows)

    for fig, cap in [
        (figs / "model_benchmark_auc_ci.png", "Supplementary Figure S1. Core benchmark AUC intervals."),
        (figs / "perturbation_bootstrap_ci.png", "Supplementary Figure S2. Perturbation confidence intervals."),
        (figs / "dag_pathway_graph.png", "Supplementary Figure S3. Directed pathway aggregation graph."),
        (figs / "external_ovarian_immune_scores.png", "Supplementary Figure S4. External ovarian immune-state score distributions."),
        (figs / "external_ovarian_immune_heatmap.png", "Supplementary Figure S5. External ovarian immune-state score correlation."),
        (figs / "immune_receptor_proxy_heatmap.png", "Supplementary Figure S6. TCGA-OV immune proxy correlation."),
        (figs / "immune_receptor_proxy_by_risk.png", "Supplementary Figure S7. TCGA-OV immune proxy by risk group."),
    ]:
        add_figure(doc, fig, cap)

    return save(doc, BASE / "JBI_Supplementary_Appendix.docx")


def build_title_page() -> Path:
    doc = Document()
    set_font(doc)
    heading(doc, "Title Page", 1)
    doc.add_paragraph(f"Title: {TITLE}")
    doc.add_paragraph(f"Short title: {SHORT_TITLE}")
    doc.add_paragraph("Author: Dr Siddalingaiah H S")
    doc.add_paragraph("Affiliation: Professor, Community Medicine, Shridevi Institute of Medical Sciences and Research Hospital, Tumkur, India")
    doc.add_paragraph("Corresponding author: Dr Siddalingaiah H S")
    doc.add_paragraph("Email: hssling@yahoo.com")
    doc.add_paragraph("Phone: 8941087719")
    doc.add_paragraph("ORCID: 0000-0002-4771-8285")
    doc.add_paragraph("Keywords: ovarian cancer; multi-omics; network biology; translational bioinformatics; reproducible research; CAR-T benchmarking")
    return save(doc, BASE / "JBI_Title_Page.docx")


def build_cover_letter() -> Path:
    doc = Document()
    set_font(doc)
    heading(doc, "Cover Letter", 1)
    doc.add_paragraph(f"Date: {date.today().isoformat()}")
    doc.add_paragraph("To,\nThe Editors\nJournal of Biomedical Informatics")
    doc.add_paragraph(f"Subject: Submission of manuscript \"{TITLE}\"")
    doc.add_paragraph(
        "This submission is directed to the Journal of Biomedical Informatics because the work is fundamentally a translational bioinformatics and reproducible pipeline study rather than a conventional single-cohort oncology report. The manuscript presents an end-to-end public-data framework for uncertainty-aware multi-omics integration in TCGA ovarian cancer, combined with external ovarian immune-context validation and a direct CAR-product raw-read benchmark."
    )
    doc.add_paragraph(
        "The informatics contribution is the explicit separation and auditing of three evidence layers: core matched TCGA-OV network inference, external ovarian immune-state validation, and direct CAR-product sequence-file benchmarking. This separation prevents overstatement while making the workflow reusable for future alignment-based and external validation extensions."
    )
    doc.add_paragraph(
        "The manuscript is original, has not been published elsewhere, and is not under consideration by another journal in this form. Previous oncology-oriented packages prepared for other journals are distinct versions and are not being reused as a concurrent submission package. This version has been reframed and rebuilt specifically for Journal of Biomedical Informatics."
    )
    doc.add_paragraph(
        "Sincerely,\n\nDr Siddalingaiah H S\nProfessor, Community Medicine\nShridevi Institute of Medical Sciences and Research Hospital, Tumkur, India\nEmail: hssling@yahoo.com\nPhone: 8941087719\nORCID: 0000-0002-4771-8285"
    )
    return save(doc, BASE / "JBI_Cover_Letter.docx")


def build_declarations() -> Path:
    doc = Document()
    set_font(doc)
    heading(doc, "Declarations", 1)
    items = [
        "Competing interests: The author declares no competing interests.",
        "Funding: This research did not receive any specific grant from funding agencies in the public, commercial, or not-for-profit sectors.",
        "Ethics: This study used only public de-identified datasets. No new institutional ethics approval or participant consent was required.",
        "Data availability: Public source data are available from GDC, GEO, ENA/SRA, cBioPortal, and PDC under their respective terms. Derived outputs and reproducibility assets are available in the project workspace.",
        "Generative AI disclosure: AI-assisted tools were used for drafting support and language refinement. The author reviewed, edited, and takes full responsibility for the manuscript content.",
    ]
    for item in items:
        doc.add_paragraph(item)
    return save(doc, BASE / "JBI_Declarations.docx")


def build_checklist() -> Path:
    doc = Document()
    set_font(doc)
    heading(doc, "JBI Submission Checklist", 1)
    items = [
        "Main manuscript prepared in Word format.",
        "Abstract is concise and under 250 words.",
        "Graphical abstract prepared as required by Journal of Biomedical Informatics guidance.",
        "Keywords included.",
        "Title page includes corresponding author details and ORCID.",
        "Declarations file prepared.",
        "Supplementary appendix prepared and cited.",
        "Generative AI disclosure included.",
        "No concurrent submission of this JBI-targeted version.",
    ]
    for item in items:
        doc.add_paragraph(f"- {item}")
    return save(doc, BASE / "JBI_Submission_Checklist.docx")


def build_readme(main_path: Path, supp_path: Path) -> None:
    text = f"""# Journal of Biomedical Informatics Submission Package

Target journal: Journal of Biomedical Informatics
Package date: {date.today().isoformat()}

## Rationale

This package reframes the project as a translational bioinformatics contribution centered on:
- reproducible multi-omics pipeline design
- uncertainty-aware network interpretation
- external ovarian immune-context validation
- direct CAR-product raw-read benchmarking

## Files

- `JBI_Main_Manuscript.docx`
- `JBI_Title_Page.docx`
- `JBI_Cover_Letter.docx`
- `JBI_Declarations.docx`
- `JBI_Supplementary_Appendix.docx`
- `JBI_Submission_Checklist.docx`
- `JBI_Graphical_Abstract.png`

## Current preferred manuscript files

- Main manuscript: `{main_path.name}`
- Supplementary appendix: `{supp_path.name}`
"""
    (BASE / "JBI_submission_bundle_readme.md").write_text(text, encoding="utf-8")


def main() -> None:
    BASE.mkdir(parents=True, exist_ok=True)
    main_path = build_main()
    supp_path = build_supplement()
    build_title_page()
    build_cover_letter()
    build_declarations()
    build_checklist()
    build_readme(main_path, supp_path)
    print(f"JBI package generated: {BASE}")


if __name__ == "__main__":
    main()
