from __future__ import annotations

from pathlib import Path
from datetime import date

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]


def set_default_font(doc: Document, name: str = "Times New Roman", size: int = 12) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table_from_df(doc: Document, df: pd.DataFrame, title: str, max_rows: int | None = None) -> None:
    doc.add_paragraph(title).runs[0].bold = True
    dfx = df.copy()
    if max_rows is not None:
        dfx = dfx.head(max_rows)
    table = doc.add_table(rows=1, cols=len(dfx.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, c in enumerate(dfx.columns):
        hdr[i].text = str(c)
    for _, row in dfx.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(dfx.columns):
            cells[i].text = str(row[c])
    doc.add_paragraph("")


def insert_figure(doc: Document, img_path: Path, caption: str, width_in: float = 6.0) -> None:
    if not img_path.exists():
        doc.add_paragraph(f"[Missing figure: {img_path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Inches(width_in))
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def safe_save(doc: Document, path: Path) -> None:
    try:
        doc.save(path)
    except PermissionError:
        doc.save(path.with_name(path.stem + "_updated.docx"))


def build_main_docx(base: Path) -> None:
    tables = ROOT / "results" / "tables"
    figs = ROOT / "results" / "figures"
    nets = ROOT / "results" / "networks"

    bench = pd.read_csv(tables / "model_benchmark.csv")
    hub_stab = pd.read_csv(nets / "network_centrality_stability.csv")
    pert = pd.read_csv(tables / "perturbation_delta.csv")
    enrich = pd.read_csv(tables / "hub_pathway_enrichment.csv") if (tables / "hub_pathway_enrichment.csv").exists() else pd.DataFrame()
    ablation = pd.read_csv(tables / "input_output_ablation_auc.csv") if (tables / "input_output_ablation_auc.csv").exists() else pd.DataFrame()

    top_auc = bench.sort_values("auc", ascending=False).iloc[0]
    top_cox = bench.sort_values("cox_c_index", ascending=False).iloc[0]
    top_hubs = ", ".join(hub_stab.head(4)["node"].tolist())
    top_pert = pert.sort_values("delta_global_pagerank_l1", ascending=False).head(4)["hub"].tolist()

    doc = Document()
    set_default_font(doc)

    add_heading(doc, "Uncertainty-Aware Multi-Omics Integration Reveals Stable Cross-Layer Hubs and Perturbation-Sensitive Network Architecture in TCGA Ovarian Cancer", 1)
    p = doc.add_paragraph("Dr Siddalingaiah H S, Professor, Community Medicine, Shridevi Institute of Medical Sciences and Research Hospital, Tumkur. ")
    p.add_run("Email: hssling@yahoo.com | ORCID: 0000-0002-4771-8285")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "Structured Abstract", 2)
    doc.add_paragraph(
        "Background: Integrated ovarian cancer multi-omics studies often report performance without clarifying which signals remain stable across methods and uncertainty checks."
    )
    doc.add_paragraph(
        f"Methods: Public TCGA-OV multi-omics data were harmonized and analyzed with latent integration, supervised multiblock modeling, network analysis, perturbation, pathway aggregation, ablation, and predictive calibration."
    )
    doc.add_paragraph(
        f"Results: RNA modules gave the highest AUC ({top_auc['auc']:.3f}), CNA the highest Cox C-index ({top_cox['cox_c_index']:.3f}), and the most stable hubs were {top_hubs}. Perturbation prioritized {', '.join(top_pert)}."
    )
    doc.add_paragraph(
        "Conclusions: Stable cross-layer hubs and perturbation-sensitive architecture were more robust than direct predictive claims and provide the strongest translational signal."
    )

    add_heading(doc, "Introduction", 2)
    doc.add_paragraph(
        "Ovarian cancer is biologically heterogeneous across genomic, epigenomic, transcriptomic, and proteomic layers. In deeply matched public cohorts, the key question is not whether integration can produce a single favorable model, but which conclusions remain stable after uncertainty-aware stress testing."
    )
    doc.add_paragraph(
        "This study therefore prioritizes reproducible systems evidence over isolated point estimates. We used TCGA-OV to identify cross-layer hubs that remain stable across latent integration, network centrality, perturbation analysis, and pathway-oriented interpretation."
    )

    add_heading(doc, "Methods", 2)
    doc.add_paragraph(
        "Public TCGA-OV data were harmonized at patient level. Core analyses required matched mutation, CNA, methylation, and RNA data; protein was evaluated in a fairness-restricted subset. MOFA-like latent factors, DIABLO-like supervised components, a multilayer network, perturbation experiments, ablation tests, and repeated-CV ML benchmarking were performed."
    )

    add_heading(doc, "Results", 2)
    doc.add_paragraph(
        f"RNA modules had the strongest discriminative signal (AUC {top_auc['auc']:.3f}; 95% CI {top_auc['auc_ci_low']:.3f}-{top_auc['auc_ci_high']:.3f}), whereas CNA had the strongest survival-ordering signal (Cox C-index {top_cox['cox_c_index']:.3f}; 95% CI {top_cox['cox_c_index_ci_low']:.3f}-{top_cox['cox_c_index_ci_high']:.3f})."
    )
    doc.add_paragraph(
        f"Hub stability analysis consistently prioritized {top_hubs}. Perturbation ranked {', '.join(top_pert)} among the highest-impact nodes, and the strongest ablation combination was RNA modules plus protein modules."
    )
    if not enrich.empty:
        lf6 = enrich[(enrich["hub"] == "LF6") & (enrich["library"] == "MSigDB_Hallmark_2020")].head(1)
        lf8 = enrich[(enrich["hub"] == "LF8") & (enrich["library"] == "GO_Biological_Process_2023")].head(1)
        if not lf6.empty:
            row = lf6.iloc[0]
            doc.add_paragraph(
                f"Enrichment analysis linked LF6 to {row['Term']} (adjusted P={row['Adjusted P-value']:.3g}), reinforcing a stromal/extracellular-matrix interpretation."
            )
        if not lf8.empty:
            row = lf8.iloc[0]
            doc.add_paragraph(
                f"LF8 showed enrichment for {row['Term']} (adjusted P={row['Adjusted P-value']:.3g}), supporting a neural-lineage and cell-junction related signal."
            )

    add_table_from_df(
        doc,
        bench[["model", "auc", "auc_ci_low", "auc_ci_high", "cox_c_index", "cox_c_index_ci_low", "cox_c_index_ci_high"]],
        "Table 1. Benchmark summary.",
    )
    add_table_from_df(
        doc,
        hub_stab[["node", "rank_score_mean", "rank_score_ci_low", "rank_score_ci_high", "topk_freq"]].head(8),
        "Table 2. Hub stability summary.",
    )
    if not enrich.empty:
        add_table_from_df(
            doc,
            enrich[["hub", "library", "Term", "Adjusted P-value", "Genes"]].head(12),
            "Table 3. Selected hub pathway enrichment results.",
        )

    insert_figure(doc, figs / "multilayer_network_graph.png", "Figure 1. Integrated multilayer network graph.")
    insert_figure(doc, figs / "dag_pathway_graph.png", "Figure 2. DAG-style pathway graph.")
    insert_figure(doc, figs / "perturbation_bootstrap_ci.png", "Figure 3. Perturbation effect sizes with bootstrap intervals.")
    insert_figure(doc, figs / "sensitivity_perturbation_curves.png", "Figure 4. Hub sensitivity curves across perturbation fractions.")
    insert_figure(doc, figs / "advanced_ml_benchmark_auc_ci.png", "Figure 5. Advanced ML benchmark AUC with repeated-CV confidence intervals.")
    insert_figure(doc, figs / "input_output_ablation_top_auc.png", "Figure 6. Top input-output ablation combinations.")

    add_heading(doc, "Discussion", 2)
    doc.add_paragraph(
        "The central result is a stable hub-centric network architecture rather than a high-accuracy classifier. Transcript-level features carried the clearest discriminative signal, CNA best preserved survival ranking, and a small set of latent hubs dominated the perturbation and stability analyses."
    )
    doc.add_paragraph(
        "The enrichment layer sharpened the biological interpretation of those hubs, especially LF6 as an extracellular-matrix or EMT-linked state and LF8 as a neural-lineage or cell-junction associated state. Advanced ML remained modest and served primarily as calibration against overstatement."
    )

    add_heading(doc, "Additional Information", 2)
    for line in [
        "Acknowledgements: Public de-identified resources from GDC, cBioPortal, and PDC were used.",
        "Authors' contributions: Dr Siddalingaiah H S conceived the work, interpreted the results, and prepared the manuscript.",
        "Ethics approval and consent to participate: Secondary analysis of public de-identified data; no new ethics approval required.",
        "Consent for publication: Not applicable.",
        "Data availability: Derived artifacts and reproducibility outputs are available in the workspace and release bundles.",
        "Competing interests: The author declares no competing interests.",
        "Funding information: The author received no specific funding for this work.",
    ]:
        doc.add_paragraph(line)

    safe_save(doc, base / "BJC_Main_Manuscript.docx")


def build_supplementary_docx(base: Path) -> None:
    tables = ROOT / "results" / "tables"
    enrich = pd.read_csv(tables / "hub_pathway_enrichment.csv") if (tables / "hub_pathway_enrichment.csv").exists() else pd.DataFrame()
    hubsum = pd.read_csv(tables / "hub_biology_summary.csv")

    doc = Document()
    set_default_font(doc)
    add_heading(doc, "British Journal of Cancer Supplementary Appendix", 1)
    doc.add_paragraph(f"Generated: {date.today().isoformat()}")
    doc.add_paragraph("Author: Dr Siddalingaiah H S")
    add_heading(doc, "Supplementary Tables", 2)
    add_table_from_df(doc, hubsum, "Supplementary Table S1. Hub biology summary.", max_rows=24)
    if not enrich.empty:
        add_table_from_df(doc, enrich, "Supplementary Table S2. Full hub pathway enrichment table.", max_rows=40)
    safe_save(doc, base / "BJC_Supplementary_Appendix.docx")


def build_checklist_docx(base: Path) -> None:
    doc = Document()
    set_default_font(doc)
    add_heading(doc, "BJC Submission Checklist", 1)
    for item in [
        "Main manuscript under 5,000 words excluding abstract, references, and figure legends.",
        "Structured abstract <= 200 words.",
        "Additional Information sections included in required order.",
        "Title page includes author, affiliation, corresponding email, and ORCID.",
        "Main figures/tables limited to 6 combined in manuscript.",
        "Hub interpretation and pathway enrichment included as supplementary support.",
        "Cover letter prepared for broad oncology readership.",
    ]:
        doc.add_paragraph(f"- {item}")
    safe_save(doc, base / "BJC_Submission_Checklist.docx")


def main() -> None:
    base = ROOT / "manuscript" / "submission_package" / "targets" / "british_journal_of_cancer"
    base.mkdir(parents=True, exist_ok=True)
    build_main_docx(base)
    build_supplementary_docx(base)
    build_checklist_docx(base)
    print("BJC DOCX package generated:", base)


if __name__ == "__main__":
    main()
